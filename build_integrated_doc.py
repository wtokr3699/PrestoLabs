from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/applw/Documents/Codex/PrestoLabs")
OUTPUT = ROOT / "PrestoLabs_통합_기획서.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_page_margins(section):
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(10.5)


doc = Document()
for section in doc.sections:
    set_page_margins(section)

styles = doc.styles
styles["Normal"].font.name = "Arial"
styles["Normal"].font.size = Pt(10.5)
styles["Normal"].paragraph_format.space_after = Pt(6)
styles["Normal"].paragraph_format.line_spacing = 1.35

styles["Title"].font.name = "Arial"
styles["Title"].font.size = Pt(24)
styles["Title"].font.bold = True
styles["Title"].font.color.rgb = RGBColor(28, 38, 58)

for name, size in [("Heading 1", 15), ("Heading 2", 12), ("Heading 3", 10.5)]:
    styles[name].font.name = "Arial"
    styles[name].font.size = Pt(size)
    styles[name].font.bold = True
    styles[name].font.color.rgb = RGBColor(28, 38, 58)

title = doc.add_paragraph(style="Title")
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
title.add_run("PrestoLabs 통합 서비스 기획서")

sub = doc.add_paragraph()
sub.paragraph_format.space_after = Pt(14)
sub_run = sub.add_run("팀 문서 통합본 | 홈페이지 및 MVP 구축 기준안")
sub_run.font.size = Pt(11)
sub_run.font.color.rgb = RGBColor(92, 103, 125)

meta = doc.add_table(rows=1, cols=3)
meta.alignment = WD_TABLE_ALIGNMENT.LEFT
meta.style = "Table Grid"
for idx, text in enumerate(["브랜드명\nPrestoLabs", "핵심 약속\n48시간 실행", "서비스 구조\nAI 기획 + 전문가 매칭"]):
    cell = meta.cell(0, idx)
    cell.text = text
    set_cell_shading(cell, "F4F2ED")

sections = [
    (
        "1. 최종 서비스 정의",
        [
            "PrestoLabs는 아이디어만 입력하면 AI가 요구사항을 정리하고, 가장 적합한 실행자를 매칭해 48시간 안에 작동하는 결과물로 연결하는 AI 기반 MVP 제작 플랫폼이다.",
            "슬로건은 '48시간 안에, 아이디어를 작동하는 결과물로.'를 채택한다.",
        ],
    ),
    (
        "2. 통합 방향",
        [
            "브랜드는 PrestoLabs로 통일한다.",
            "거래 신뢰 구조는 WorkBridge의 AI 요구사항 정리, 계약, 에스크로, 검수 흐름을 반영한다.",
            "차별화 메시지는 뚝딱연구소의 48시간 MVP 실행 모델을 채택한다.",
            "사용자 경험은 BRIDGE의 AI 기획 번역 중심 플로우를 반영한다.",
        ],
    ),
    (
        "3. 문제 정의",
        [
            "의뢰인은 무엇을 어떻게 맡겨야 하는지 모르고, 견적과 결과물의 품질을 신뢰하기 어렵다.",
            "제작자는 AI를 활용한 빠른 제작 역량을 정당한 가격으로 인정받기 어렵다.",
            "기존 플랫폼은 검색과 견적 비교는 제공하지만, 범위 합의와 검수 기준의 표준화가 약하다.",
        ],
    ),
    (
        "4. 핵심 가치 제안",
        [
            "아이디어만 있어도 시작 가능",
            "AI가 요구사항을 구조화",
            "48시간 안에 작동하는 결과물 전달",
            "에스크로와 체크리스트 기반의 안전한 거래",
        ],
    ),
    (
        "5. 서비스 구조",
        [
            "트랙 A. 빠른 패키지형: 랜딩페이지, 예약 시스템, 업무 자동화처럼 반복 수요를 표준 상품으로 제공한다.",
            "트랙 B. 커스텀 매칭형: 자연어 아이디어를 AI가 정리하고, 전문가 또는 스쿼드를 추천한다.",
        ],
    ),
    (
        "6. 핵심 기능",
        [
            "AI 기획 번역기",
            "AI 전문가 매칭",
            "프로젝트 등록 및 비교",
            "전자 계약 및 에스크로 결제",
            "검수 체크리스트",
            "48시간 카운트다운 대시보드",
        ],
    ),
    (
        "7. 사용자 여정",
        [
            "의뢰인: 아이디어 입력 → AI 분석 → 실행자 선택 → 결제 → 진행 확인 → 검수 → 리뷰",
            "실행자: 프로필 등록 → 추천 프로젝트 확인 → 참여 → 작업 → 제출 → 정산",
        ],
    ),
    (
        "8. MVP 범위",
        [
            "메인 랜딩 페이지",
            "아이디어 입력 인터페이스",
            "AI 분석 결과 화면",
            "추천 실행자/패키지 카드",
            "프로젝트 진행 화면",
            "결과물 제출 및 승인 구조",
        ],
    ),
]

for heading, bullets in sections:
    doc.add_paragraph(heading, style="Heading 1")
    for item in bullets:
        add_bullet(doc, item)

doc.add_paragraph("9. 차별점 비교", style="Heading 1")
table = doc.add_table(rows=1, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = "Table Grid"
hdr = table.rows[0].cells
hdr[0].text = "항목"
hdr[1].text = "기존 플랫폼"
hdr[2].text = "PrestoLabs"
for cell in hdr:
    set_cell_shading(cell, "E6ECF7")

rows = [
    ("시작 방식", "전문가를 직접 검색", "아이디어만 입력"),
    ("기획 지원", "직접 정리", "AI가 요구사항 자동 정리"),
    ("매칭 구조", "검색/저가 비교", "AI 추천 중심"),
    ("진행 체감", "메시지 협의 중심", "상태 표시 + 카운트다운"),
    ("거래 신뢰", "플랫폼별 편차", "에스크로 + 체크리스트"),
]
for a, b, c in rows:
    row = table.add_row().cells
    row[0].text = a
    row[1].text = b
    row[2].text = c

doc.add_paragraph("10. 디자인 원칙", style="Heading 1")
for item in [
    "김태민 시안의 미니멀 카드형 편집 레이아웃을 반영한다.",
    "뉴트럴 배경, 얇은 보더, 절제된 강조색, 넉넉한 여백을 사용한다.",
    "서비스 흐름이 보이는 전환형 랜딩 구조를 채택한다.",
]:
    add_bullet(doc, item)

doc.add_section(WD_SECTION_START.NEW_PAGE)
doc.add_paragraph("부록. 메인 페이지 메시지 구조", style="Heading 1")
for item in [
    "Hero: 아이디어만 적어주세요. PrestoLabs가 AI로 기획을 정리하고 48시간 안에 실행합니다.",
    "핵심 섹션: 아이디어 입력, AI 기획 정리, 전문가 매칭, 안전한 결제, 진행 관리, 결과물 전달",
    "신뢰 섹션: 체크리스트 검수, 명확한 범위 합의, 진행 상태 가시화, 빠른 납기",
]:
    add_bullet(doc, item)

doc.save(OUTPUT)
print(OUTPUT)
